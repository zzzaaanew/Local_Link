"""취재 계획서 → 문서 파일 (docx / hwpx / pdf).

입력은 weekly_plan 프롬프트가 만든 마크다운(plan_markdown)이다.
세 형식 모두 같은 파싱 결과(_parse)를 공유하고, 출력 단계만 다르다.

  docx  python-docx. 한글(한컴오피스) 2010 이상에서 바로 열린다.
  hwpx  한글 문서의 개방형 표준(OWPML, KS X 6101). ZIP+XML 구조를 직접 만든다.
  pdf   reportlab. 한국어는 내장 CID 폰트(HYSMyeongJo-Medium)를 쓴다.

파일명은 취재 제안서 제목을 그대로 쓴다.
"""
import io
import re
import zipfile
from datetime import datetime
from xml.sax.saxutils import escape

FONT = "맑은 고딕"
FORMATS = ("docx", "hwpx", "pdf")

MEDIA_TYPES = {
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "hwpx": "application/hwp+zip",
    "pdf": "application/pdf",
}


def build(pitch: dict, fmt: str = "docx") -> tuple[bytes, str, str]:
    """(파일 바이트, 파일명, media type)."""
    if fmt not in FORMATS:
        raise ValueError(f"지원하지 않는 형식입니다: {fmt}")

    title = (pitch.get("title") or "취재 계획서").strip()
    blocks = _parse(pitch)
    builder = {"docx": _build_docx, "hwpx": _build_hwpx, "pdf": _build_pdf}[fmt]
    return builder(title, blocks), f"{_safe_name(title)}.{fmt}", MEDIA_TYPES[fmt]


# ------------------------------------------------------------
# 공통 파싱 — 마크다운을 (종류, 텍스트) 블록 목록으로
# 종류: title | meta | stamp | h1 | h2 | li | p | gap
# ------------------------------------------------------------

def _parse(pitch: dict) -> list[tuple[str, str]]:
    blocks: list[tuple[str, str]] = [("title", (pitch.get("title") or "취재 계획서").strip())]

    meta = " · ".join(
        v
        for v in [
            pitch.get("agency"),
            pitch.get("region"),
            pitch.get("timing"),
            f"발표 예정 {pitch.get('scheduled_at')}" if pitch.get("scheduled_at") else "",
        ]
        if v
    )
    if meta:
        blocks.append(("meta", meta))
    blocks.append(("stamp", f"생성일 {datetime.now().strftime('%Y-%m-%d')} · LocalDeskAI"))
    blocks.append(("gap", ""))

    body = (pitch.get("plan_markdown") or "").strip()
    if body:
        blocks.extend(_parse_markdown(body))
    else:
        for label, key in (("발제 이유", "why"), ("취재 포인트", "summary")):
            value = (pitch.get(key) or "").strip()
            if value:
                blocks.append(("h1", label))
                blocks.append(("p", value))
    return blocks


def _parse_markdown(text: str) -> list[tuple[str, str]]:
    out: list[tuple[str, str]] = []
    blank = False  # 빈 줄이 연달아 나와도 여백은 한 줄만
    for raw in text.replace("\r\n", "\n").split("\n"):
        line = raw.strip()

        if not line or (set(line) <= {"-", "─", "━", "=", "*", "_"} and len(line) >= 3):
            if not blank and out:
                out.append(("gap", ""))
                blank = True
            continue
        blank = False

        heading = re.match(r"^(#{1,6})\s*(.+)$", line)
        if heading:
            kind = "h1" if len(heading.group(1)) <= 2 else "h2"
            out.append((kind, _clean(heading.group(2))))
            continue

        bullet = re.match(r"^[-*•]\s+(.+)$", line)
        if bullet:
            out.append(("li", _clean(bullet.group(1))))
            continue

        numbered = re.match(r"^\d+[.)]\s+(.+)$", line)
        if numbered:
            out.append(("li", _clean(numbered.group(1))))
            continue

        out.append(("p", _clean(line)))
    return out


def _clean(text: str) -> str:
    """굵게/기울임 표시가 문서에 기호로 남지 않게 제거한다."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", text)
    return text.replace("`", "")


def _safe_name(title: str) -> str:
    """파일명으로 못 쓰는 문자를 제거한다. 대괄호 제목도 그대로 살린다."""
    safe = re.sub(r'[\\/:*?"<>|\n\r\t]', "", title).strip()
    safe = re.sub(r"\s+", " ", safe)
    return safe[:80] or "취재계획서"


# ------------------------------------------------------------
# docx
# ------------------------------------------------------------

def _build_docx(title: str, blocks: list[tuple[str, str]]) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt

    def style(run, size: float, bold: bool = False) -> None:
        run.font.name = FONT
        run.font.size = Pt(size)
        run.bold = bold
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    for kind, text in blocks:
        if kind == "gap":
            doc.add_paragraph()
            continue
        if kind == "li":
            para = doc.add_paragraph(text, style="List Bullet")
            style(para.runs[0], 10.5)
            continue

        para = doc.add_paragraph(text)
        if kind == "title":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style(para.runs[0], 18, bold=True)
        elif kind == "meta":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style(para.runs[0], 9.5)
        elif kind == "stamp":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style(para.runs[0], 8.5)
        elif kind == "h1":
            style(para.runs[0], 13, bold=True)
        elif kind == "h2":
            style(para.runs[0], 11.5, bold=True)
        else:
            style(para.runs[0], 10.5)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ------------------------------------------------------------
# pdf
# ------------------------------------------------------------

def _build_pdf(title: str, blocks: list[tuple[str, str]]) -> bytes:
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    # 한국어 내장 CID 폰트. 외부 폰트 파일 없이 한글이 렌더링된다.
    for name in ("HYSMyeongJo-Medium", "HYGothic-Medium"):
        try:
            pdfmetrics.registerFont(UnicodeCIDFont(name))
        except Exception:  # 이미 등록된 경우 등
            pass

    body_font, head_font = "HYSMyeongJo-Medium", "HYGothic-Medium"
    styles = {
        "title": ParagraphStyle("t", fontName=head_font, fontSize=17, leading=24, alignment=TA_CENTER, spaceAfter=6),
        "meta": ParagraphStyle("m", fontName=body_font, fontSize=9, leading=14, alignment=TA_CENTER, textColor="#555555"),
        "stamp": ParagraphStyle("s", fontName=body_font, fontSize=8, leading=13, alignment=TA_CENTER, textColor="#888888", spaceAfter=10),
        "h1": ParagraphStyle("h1", fontName=head_font, fontSize=12.5, leading=19, spaceBefore=12, spaceAfter=4),
        "h2": ParagraphStyle("h2", fontName=head_font, fontSize=11, leading=17, spaceBefore=8, spaceAfter=3),
        "p": ParagraphStyle("p", fontName=body_font, fontSize=10, leading=17),
        "li": ParagraphStyle("li", fontName=body_font, fontSize=10, leading=17, leftIndent=12, bulletIndent=2),
    }

    story = []
    for kind, text in blocks:
        if kind == "gap":
            story.append(Spacer(1, 5))
            continue
        if kind == "li":
            # ListFlowable 대신 bulletText 를 쓴다. 글머리 기호도 본문 폰트로 그려지고,
            # 항목이 누락될 여지가 없다.
            story.append(Paragraph(_pdf_text(text), styles["li"], bulletText="-"))
            continue
        story.append(Paragraph(_pdf_text(text), styles.get(kind, styles["p"])))

    buffer = io.BytesIO()
    SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=22 * mm, rightMargin=22 * mm,
        topMargin=20 * mm, bottomMargin=20 * mm,
        title=title, author="LocalDeskAI",
    ).build(story)
    return buffer.getvalue()


# 내장 CID 폰트(HYSMyeongJo-Medium)에 글리프가 없어 화면에서 사라지는 문자들.
# 빠진 자리는 티가 나지 않아 조용히 문장이 깨지므로 미리 안전한 문자로 바꾼다.
# docx·hwpx 는 시스템 폰트를 쓰므로 이 치환을 적용하지 않는다.
_PDF_SAFE = str.maketrans({
    "·": "|", "∙": "-", "•": "-", "‧": "-", "ㆍ": "-",
    "—": "-", "–": "-", "―": "-",
    "✓": "V", "✔": "V", "☑": "V",
    "…": "...",
})


def _pdf_text(text: str) -> str:
    """PDF용 텍스트. 글리프 치환 후, Paragraph 의 미니 HTML 파서를 위해 &<> 이스케이프."""
    return escape(text.translate(_PDF_SAFE))


# ------------------------------------------------------------
# hwpx — 한글 문서 개방형 표준(OWPML). ZIP + XML 컨테이너.
# ------------------------------------------------------------

_HWPX_VERSION = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<hv:HCFVersion xmlns:hv="http://www.hancom.co.kr/hwpml/2011/version"'
    ' tagetApplication="WORDPROCESSOR" major="5" minor="0" micro="5" buildNumber="0"'
    ' os="1" application="LocalDeskAI"/>'
)

_HWPX_CONTAINER = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<ocf:container xmlns:ocf="urn:oasis:names:tc:opendocument:xmlns:container"'
    ' xmlns="http://www.idpf.org/2007/opf/">'
    '<ocf:rootfiles>'
    '<ocf:rootfile full-path="Contents/content.hpf" media-type="application/hwpml-package+xml"/>'
    '</ocf:rootfiles></ocf:container>'
)

_HWPX_MANIFEST = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<odf:manifest xmlns:odf="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0" odf:version="1.2">'
    '<odf:file-entry odf:full-path="/" odf:media-type="application/hwp+zip"/>'
    '</odf:manifest>'
)

_HWPX_SETTINGS = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<ha:HWPApplicationSetting xmlns:ha="http://www.hancom.co.kr/hwpml/2011/app"/>'
)

# 문단 모양(ParaShape) / 글자 모양(CharShape) 인덱스
#   0 본문 · 1 제목(가운데, 큰 글씨) · 2 메타(가운데, 작게) · 3 소제목 · 4 목록
_HWPX_CONTENT_HPF = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<opf:package xmlns:opf="http://www.idpf.org/2007/opf/" version=""'
    ' unique-identifier="" id="">'
    '<opf:metadata xmlns:opf="http://www.idpf.org/2007/opf/"'
    ' xmlns:dc="http://purl.org/dc/elements/1.1/">'
    '<opf:title>{title}</opf:title>'
    '<opf:language>ko</opf:language>'
    '<opf:meta name="creator" content="LocalDeskAI"/>'
    '</opf:metadata>'
    '<opf:manifest>'
    '<opf:item id="header" href="Contents/header.xml" media-type="application/xml"/>'
    '<opf:item id="section0" href="Contents/section0.xml" media-type="application/xml"/>'
    '<opf:item id="settings" href="settings.xml" media-type="application/xml"/>'
    '</opf:manifest>'
    '<opf:spine><opf:itemref idref="section0" linear="yes"/></opf:spine>'
    '</opf:package>'
)


def _build_hwpx(title: str, blocks: list[tuple[str, str]]) -> bytes:
    buffer = io.BytesIO()
    # mimetype 은 압축하지 않고 가장 먼저 넣는다 (OPC 규약).
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(
            zipfile.ZipInfo("mimetype"), "application/hwp+zip", compress_type=zipfile.ZIP_STORED
        )
        zf.writestr("version.xml", _HWPX_VERSION)
        zf.writestr("META-INF/container.xml", _HWPX_CONTAINER)
        zf.writestr("META-INF/manifest.xml", _HWPX_MANIFEST)
        zf.writestr("settings.xml", _HWPX_SETTINGS)
        zf.writestr("Contents/content.hpf", _HWPX_CONTENT_HPF.format(title=escape(title)))
        zf.writestr("Contents/header.xml", _hwpx_header())
        zf.writestr("Contents/section0.xml", _hwpx_section(blocks))
        zf.writestr("Preview/PrvText.txt", _hwpx_preview(blocks))
    return buffer.getvalue()


def _hwpx_header() -> str:
    """글자 모양·문단 모양 정의. 크기 단위는 HWPUNIT(1pt = 100)."""
    char_shapes = [
        (1050, 0),   # 0 본문
        (1800, 1),   # 1 제목 (굵게)
        (950, 0),    # 2 메타
        (1300, 1),   # 3 소제목 큰 것
        (1150, 1),   # 4 소제목 작은 것
        (850, 0),    # 5 생성일
    ]
    char_xml = "".join(
        f'<hh:charPr id="{i}" height="{h}" textColor="#000000" shadeColor="none"'
        f' useFontSpace="0" useKerning="0">'
        f'<hh:fontRef hangul="0" latin="0" hanja="0" japanese="0" other="0" symbol="0" user="0"/>'
        f'<hh:ratio hangul="100" latin="100" hanja="100" japanese="100" other="100" symbol="100" user="100"/>'
        f'<hh:spacing hangul="0" latin="0" hanja="0" japanese="0" other="0" symbol="0" user="0"/>'
        f'<hh:relSz hangul="100" latin="100" hanja="100" japanese="100" other="100" symbol="100" user="100"/>'
        f'<hh:offset hangul="0" latin="0" hanja="0" japanese="0" other="0" symbol="0" user="0"/>'
        f'<hh:bold>{b}</hh:bold>'
        f'</hh:charPr>'
        for i, (h, b) in enumerate(char_shapes)
    )

    # (정렬, 왼쪽 여백, 위 여백, 아래 여백)
    para_shapes = [
        ("JUSTIFY", 0, 0, 0),      # 0 본문
        ("CENTER", 0, 0, 300),     # 1 제목
        ("CENTER", 0, 0, 100),     # 2 메타
        ("LEFT", 0, 600, 200),     # 3 소제목
        ("LEFT", 800, 0, 0),       # 4 목록
    ]
    para_xml = "".join(
        f'<hh:paraPr id="{i}" tabPrIDRef="0" condense="0" fontLineHeight="0"'
        f' snapToGrid="1" suppressLineNumbers="0" checked="0">'
        f'<hh:align horizontal="{align}" vertical="BASELINE"/>'
        f'<hh:heading type="NONE" idRef="0" level="0"/>'
        f'<hh:breakSetting breakLatinWord="KEEP_WORD" breakNonLatinWord="KEEP_WORD"'
        f' widowOrphan="0" keepWithNext="0" keepLines="0" pageBreakBefore="0" lineWrap="BREAK"/>'
        f'<hh:margin>'
        f'<hc:intent xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" value="0" unit="HWPUNIT"/>'
        f'<hc:left xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" value="{left}" unit="HWPUNIT"/>'
        f'<hc:right xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" value="0" unit="HWPUNIT"/>'
        f'<hc:prev xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" value="{before}" unit="HWPUNIT"/>'
        f'<hc:next xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" value="{after}" unit="HWPUNIT"/>'
        f'</hh:margin>'
        f'<hh:lineSpacing type="PERCENT" value="160" unit="HWPUNIT"/>'
        f'<hh:border borderFillIDRef="1" offsetLeft="0" offsetRight="0" offsetTop="0"'
        f' offsetBottom="0" connect="0" ignoreMargin="0"/>'
        f'</hh:paraPr>'
        for i, (align, left, before, after) in enumerate(para_shapes)
    )

    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head"'
        ' xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" version="1.4" secCnt="1">'
        '<hh:beginNum page="1" footnote="1" endnote="1" pic="1" tbl="1" equation="1"/>'
        '<hh:refList>'
        '<hh:fontfaces itemCnt="1">'
        '<hh:fontface lang="HANGUL" fontCnt="1">'
        f'<hh:font id="0" face="{FONT}" type="TTF" isEmbedded="0">'
        '<hh:typeInfo familyType="FCAT_GOTHIC" weight="8" proportion="4" contrast="0"'
        ' strokeVariation="1" armStyle="1" letterform="1" midline="1" xHeight="1"/>'
        '</hh:font></hh:fontface>'
        '</hh:fontfaces>'
        '<hh:borderFills itemCnt="2">'
        '<hh:borderFill id="1" threeD="0" shadow="0" centerLine="NONE" breakCellSeparateLine="0">'
        '<hh:slash type="NONE" Crooked="0" isCounter="0"/>'
        '<hh:backSlash type="NONE" Crooked="0" isCounter="0"/>'
        '<hh:leftBorder type="NONE" width="0.1 mm" color="#000000"/>'
        '<hh:rightBorder type="NONE" width="0.1 mm" color="#000000"/>'
        '<hh:topBorder type="NONE" width="0.1 mm" color="#000000"/>'
        '<hh:bottomBorder type="NONE" width="0.1 mm" color="#000000"/>'
        '<hh:diagonal type="SOLID" width="0.1 mm" color="#000000"/>'
        '</hh:borderFill>'
        '</hh:borderFills>'
        f'<hh:charProperties itemCnt="{len(char_shapes)}">{char_xml}</hh:charProperties>'
        '<hh:tabProperties itemCnt="1">'
        '<hh:tabPr id="0" autoTabLeft="0" autoTabRight="0"/>'
        '</hh:tabProperties>'
        f'<hh:paraProperties itemCnt="{len(para_shapes)}">{para_xml}</hh:paraProperties>'
        '<hh:styles itemCnt="1">'
        '<hh:style id="0" type="PARA" name="바탕글" engName="Normal" paraPrIDRef="0"'
        ' charPrIDRef="0" nextStyleIDRef="0" langID="1042" lockForm="0"/>'
        '</hh:styles>'
        '</hh:refList>'
        '</hh:head>'
    )


# 블록 종류 → (문단모양 id, 글자모양 id)
_HWPX_SHAPE = {
    "title": (1, 1),
    "meta": (2, 2),
    "stamp": (2, 5),
    "h1": (3, 3),
    "h2": (3, 4),
    "li": (4, 0),
    "p": (0, 0),
    "gap": (0, 0),
}


def _hwpx_section(blocks: list[tuple[str, str]]) -> str:
    paragraphs = []
    for index, (kind, text) in enumerate(blocks):
        para_id, char_id = _HWPX_SHAPE.get(kind, (0, 0))
        content = "" if kind == "gap" else text
        if kind == "li":
            content = f"· {content}"
        paragraphs.append(
            f'<hp:p id="{index}" paraPrIDRef="{para_id}" styleIDRef="0" pageBreak="0"'
            f' columnBreak="0" merged="0">'
            f'<hp:run charPrIDRef="{char_id}">'
            f'<hp:t>{escape(content)}</hp:t>'
            f'</hp:run>'
            f'<hp:linesegarray>'
            f'<hp:lineseg textpos="0" vertpos="0" vertsize="1000" textheight="1000"'
            f' baseline="850" spacing="600" horzpos="0" horzsize="42520" flags="393216"/>'
            f'</hp:linesegarray>'
            f'</hp:p>'
        )

    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section"'
        ' xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph"'
        ' xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core">'
        + "".join(paragraphs)
        + '</hs:sec>'
    )


def _hwpx_preview(blocks: list[tuple[str, str]]) -> str:
    """한글 파일 탐색기 미리보기용 평문."""
    lines = [text for kind, text in blocks if kind != "gap" and text]
    return "\n".join(lines)[:1000]
